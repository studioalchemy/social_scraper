"""Render the analyzer's JSON report as a formal .docx file (in-memory)."""
import logging
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


# ── Styling primitives ────────────────────────────────────────────────────────

BODY_FONT = "Calibri"
HEADING_FONT = "Calibri"
INK_DARK = RGBColor(0x1f, 0x1f, 0x1f)
INK_MUTED = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0x2c, 0x3e, 0x50)  # deep slate for headings


def _set_cell_shading(cell, hex_fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _configure_styles(doc: Document) -> None:
    """Set document-wide defaults: 1-inch margins, Calibri body, larger spacing."""
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK_DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for lvl, size, color in [
        ("Heading 1", 22, ACCENT),
        ("Heading 2", 16, ACCENT),
        ("Heading 3", 13, INK_DARK),
        ("Heading 4", 11, INK_DARK),
    ]:
        st = doc.styles[lvl]
        st.font.name = HEADING_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(16 if lvl == "Heading 1" else 12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True


# ── Atomic add-* helpers ──────────────────────────────────────────────────────

def _add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = HEADING_FONT
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    p.paragraph_format.space_after = Pt(4)


def _add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = INK_MUTED
    p.paragraph_format.space_after = Pt(18)


def _add_meta_line(doc, text, *, center=False, italic=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.italic = italic
    run.font.color.rgb = INK_MUTED


def _add_italic_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.italic = True
    run.font.color.rgb = INK_MUTED


def _add_para(doc, text, *, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = bold


def _add_bullets(doc, items):
    for item in items or []:
        if not item:
            continue
        doc.add_paragraph(str(item), style="List Bullet")


def _set_col_widths(table, widths_inches: list[float]) -> None:
    """Apply per-column widths in inches. Word ignores autofit-on widths and
    sometimes ignores column-level widths, so we set width on every cell —
    that's the only setting that sticks reliably across Word builds."""
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def _add_table(doc, headers, rows, col_widths: list[float] | None = None):
    if not rows:
        _add_para(doc, "(no data)")
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        _set_cell_shading(hdr[i], "2C3E50")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, key in enumerate(headers):
            val = row.get(key) if isinstance(row, dict) else (row[c_idx] if c_idx < len(row) else "")
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(_fmt_cell(val))
            run.font.size = Pt(10)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths and len(col_widths) == len(headers):
        _set_col_widths(table, col_widths)


def _fmt_cell(v) -> str:
    if v is None or v == "" or v == 0 or v == "0":
        return "—"
    return str(v)


def _section_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.add_run("")


# ── Section renderers ────────────────────────────────────────────────────────

def _render_header(doc, report):
    header = report.get("header", {})
    _add_title(doc, header.get("title_line", "Instagram Strategy"))
    _add_subtitle(doc, header.get("subtitle", "Competitive Analysis & Recommendations"))
    _add_meta_line(doc, header.get("prepared_for", ""), center=True)
    _add_meta_line(doc, header.get("data_note", ""), center=True, italic=True)
    doc.add_paragraph()

    doc.add_heading("ACCOUNTS ANALYSED", level=3)
    _add_para(doc, header.get("accounts_analysed_line", ""))

    doc.add_heading("ANCHORED TO BUSINESS PROBLEMS", level=3)
    for bp in header.get("anchored_to_bps", []) or []:
        _add_para(doc, bp)

    doc.add_page_break()


def _render_section1(doc, section1):
    doc.add_heading("SECTION 1 — COMPETITIVE LANDSCAPE", level=1)
    if section1.get("intro_italic"):
        _add_italic_para(doc, section1["intro_italic"])

    doc.add_heading("Competitive Landscape at a Glance", level=2)
    _add_table(
        doc,
        ["account", "followers", "following", "posts", "verified", "avg_er", "format_mix", "posts_per_week"],
        section1.get("at_a_glance", []),
    )
    for fn in section1.get("footnotes", []) or []:
        _add_meta_line(doc, fn, italic=True)

    for i, acc in enumerate(section1.get("accounts", []) or [], start=1):
        doc.add_heading(f"1.{i} {acc.get('name', '')} ({acc.get('handle', '')})", level=2)
        _add_meta_line(doc, acc.get("header_line", ""))

        if acc.get("is_limited_data"):
            doc.add_heading("Key Observations", level=3)
            _add_bullets(doc, acc.get("key_observations", []))
            doc.add_heading("Key Insight", level=3)
            _add_para(doc, acc.get("key_insight", ""))
            continue

        doc.add_heading("Content Format Mix", level=3)
        _add_bullets(doc, acc.get("content_format_mix", []))

        doc.add_heading("Posting Frequency", level=3)
        _add_para(doc, acc.get("posting_frequency", ""))

        doc.add_heading("Top 5 Performing Posts", level=3)
        # Usable width on US Letter w/ 1" margins ≈ 6.5". Spend the budget on
        # the two prose columns so "why_it_worked" doesn't get squeezed to a
        # 0.8" sliver that wraps every other word.
        _add_table(
            doc,
            ["rank", "caption_excerpt", "type", "likes", "comments", "why_it_worked"],
            acc.get("top_5_posts", []),
            col_widths=[0.45, 1.5, 0.55, 0.65, 0.75, 2.6],
        )

        doc.add_heading("Content Themes & Pillars", level=3)
        _add_bullets(doc, acc.get("content_themes", []))

        doc.add_heading("Brand Voice", level=3)
        _add_bullets(doc, acc.get("brand_voice", []))

        if acc.get("engagement_patterns"):
            doc.add_heading("Engagement Patterns", level=3)
            _add_bullets(doc, acc.get("engagement_patterns", []))

        collabs = acc.get("creator_collaborations") or []
        if collabs:
            doc.add_heading("Creator Collaborations", level=3)
            _add_table(
                doc,
                ["creator_handle", "post_type", "post_url", "caption_excerpt", "engagement_summary", "is_paid_partnership"],
                collabs,
            )
        else:
            doc.add_heading("Creator Collaborations", level=3)
            _add_meta_line(doc, "No creator collaborations detected in this scrape window.", italic=True)

        doc.add_heading("Key Insight", level=3)
        _add_para(doc, acc.get("key_insight", ""))

    doc.add_page_break()


def _render_section2(doc, section2):
    doc.add_heading("SECTION 2 — COMPETITIVE POSITIONING MAP", level=1)
    _add_para(doc, section2.get("axis1", ""))
    _add_para(doc, section2.get("axis2", ""))

    _add_table(
        doc,
        ["brand", "axis1_position", "axis2_position", "content_territory"],
        section2.get("positioning_table", []),
    )

    doc.add_heading("Strategic Positioning Summary", level=2)
    _add_para(doc, section2.get("strategic_positioning_summary", ""))

    doc.add_heading(section2.get("critical_observation_title", "Critical Observation"), level=2)
    _add_para(doc, section2.get("critical_observation_body", ""))

    doc.add_page_break()


def _render_section3(doc, section3):
    doc.add_heading("SECTION 3 — OWN ACCOUNT AUDIT", level=1)
    _add_meta_line(doc, section3.get("header_line", ""))

    doc.add_heading("3.1 What the Brand Is Doing Well", level=2)
    _add_bullets(doc, section3.get("doing_well", []))

    doc.add_heading("3.2 What the Brand Is Doing Poorly (In Isolation)", level=2)
    _add_bullets(doc, section3.get("doing_poorly_isolated", []))

    doc.add_heading("3.3 What the Brand Is Doing Poorly Relative to Competition", level=2)
    _add_bullets(doc, section3.get("doing_poorly_vs_competition", []))

    doc.add_heading("3.4 Engagement Rate vs Competitive Set", level=2)
    _add_para(doc, section3.get("er_comparison", ""))
    if section3.get("status_summary_bold"):
        _add_para(doc, section3["status_summary_bold"], bold=True)

    doc.add_heading("3.5 Content Calendar Pattern", level=2)
    _add_bullets(doc, section3.get("content_calendar_pattern", []))

    doc.add_heading("3.6 Creator & Collaboration Presence", level=2)
    _add_bullets(doc, section3.get("creator_collaboration_presence", []))

    doc.add_page_break()


def _render_section4(doc, section4):
    doc.add_heading("SECTION 4 — WHITESPACE ANALYSIS", level=1)

    for title, key in [
        ("4.1 Content Themes No Brand Is Owning", "themes_no_brand_owns"),
        ("4.2 Formats Being Underutilised Across the Category", "underutilised_formats"),
        ("4.3 Emotional Territories Available", "emotional_territories"),
        ("4.4 Audience Conversations Not Being Picked Up", "audience_conversations"),
        ("4.5 Visual & Aesthetic Whitespace", "visual_whitespace"),
    ]:
        doc.add_heading(title, level=2)
        _add_bullets(doc, section4.get(key, []))

    doc.add_page_break()


def _render_section5(doc, section5):
    doc.add_heading("SECTION 5 — RECOMMENDATIONS ANCHORED TO BUSINESS PROBLEMS", level=1)
    for block in section5.get("blocks", []) or []:
        letter = block.get("letter", "?")
        prob = block.get("problem", "")
        doc.add_heading(f"5{letter} — SOLVING {prob.upper()}", level=2)

        doc.add_heading(f"5{letter}.1 What Competitors Are Doing Successfully", level=3)
        _add_bullets(doc, block.get("competitors_doing_successfully", []))

        doc.add_heading(f"5{letter}.2 Underrepresented Occasions / Moments", level=3)
        _add_bullets(doc, block.get("underrepresented_occasions", []))

        doc.add_heading(f"5{letter}.3 Instagram Content Strategies", level=3)
        _add_bullets(doc, block.get("content_strategies", []))

        doc.add_heading(f"5{letter}.4 Recommended Content Territories", level=3)
        for terr in block.get("recommended_territories", []) or []:
            _add_para(doc, terr.get("territory_name", ""), bold=True)
            _add_para(doc, f"Anchor: {terr.get('anchor', '')}")
            _add_para(doc, f"Caption example: {terr.get('caption_example', '')}")
            _add_para(doc, f"Why it solves the BP: {terr.get('why_solves_bp', '')}")
            doc.add_paragraph()

    doc.add_page_break()


def _render_section6(doc, section6):
    doc.add_heading("SECTION 6 — CONTENT STRATEGY BLUEPRINT", level=1)

    doc.add_heading("6.1 Content Pillars", level=2)
    _add_table(
        doc,
        ["pillar_name", "content", "bp_addressed", "engagement_potential"],
        section6.get("content_pillars", []),
    )

    doc.add_heading("6.2 Format Mix Recommendation", level=2)
    _add_table(
        doc,
        ["format", "pct_of_mix", "posts_per_week", "story_usage", "rationale"],
        section6.get("format_mix", []),
    )

    doc.add_heading("6.3 Posting Cadence", level=2)
    _add_bullets(doc, section6.get("posting_cadence", []))

    doc.add_heading("6.4 Tone of Voice Direction", level=2)
    tov = section6.get("tone_of_voice", {}) or {}

    doc.add_heading("Defining Characteristics", level=3)
    for ch in tov.get("defining_characteristics", []) or []:
        _add_para(doc, ch.get("label", ""), bold=True)
        _add_para(doc, ch.get("explanation", ""))

    doc.add_heading("Example Caption Styles", level=3)
    _add_bullets(doc, tov.get("example_captions", []))

    doc.add_heading("Clear Don'ts — Based on Competitive Learnings", level=3)
    _add_bullets(doc, tov.get("donts", []))

    doc.add_heading("6.5 Creator Strategy", level=2)
    cs = section6.get("creator_strategy", {}) or {}
    _add_table(
        doc,
        ["archetype", "size", "content_genre", "platforms", "fit_for_brand"],
        cs.get("archetype_table", []),
    )

    for label_key, layer_key in [
        ("ALWAYS-ON LAYER", "always_on_layer"),
        ("CAMPAIGN LAYER", "campaign_layer"),
    ]:
        layer = cs.get(layer_key, {}) or {}
        doc.add_heading(label_key, level=3)
        _add_para(doc, f"Brief: {layer.get('brief', '—')}")
        _add_para(doc, f"Budget: {layer.get('budget', '—')}")
        _add_para(doc, f"Content approach: {layer.get('content_approach', '—')}")

    add_layer = cs.get("additional_layer", {}) or {}
    if add_layer:
        doc.add_heading(add_layer.get("name", "ADDITIONAL LAYER"), level=3)
        _add_para(doc, f"Brief: {add_layer.get('brief', '—')}")
        _add_para(doc, f"Budget: {add_layer.get('budget', '—')}")
        _add_para(doc, f"Content approach: {add_layer.get('content_approach', '—')}")

    doc.add_heading("6.6 Craft Standards", level=2)
    cs2 = section6.get("craft_standards", {}) or {}
    doc.add_heading("Hook Style & Duration for Reels", level=3)
    _add_bullets(doc, cs2.get("hook_style", []))
    doc.add_heading("Visual Treatment", level=3)
    _add_bullets(doc, cs2.get("visual_treatment", []))
    doc.add_heading("Caption Structure Guidance", level=3)
    _add_bullets(doc, cs2.get("caption_structure", []))

    doc.add_page_break()


def _render_summary(doc, summary):
    doc.add_heading("SUMMARY — TOP 10 ACTIONABLE RECOMMENDATIONS", level=1)
    if summary.get("intro_italic"):
        _add_italic_para(doc, summary["intro_italic"])

    for title, key in [
        ("IMMEDIATE (0–4 weeks)", "immediate"),
        ("SHORT TERM (1–3 months)", "short_term"),
        ("MEDIUM TERM (3–6 months)", "medium_term"),
    ]:
        doc.add_heading(title, level=2)
        for rec in summary.get(key, []) or []:
            _add_para(doc, rec.get("label", ""), bold=True)
            _add_para(doc, rec.get("action", ""))
            _add_meta_line(doc, f"Impacts: {rec.get('bp_impact', '—')}", italic=True)
            _add_meta_line(doc, f"Effort: {rec.get('effort_signal', '—')}", italic=True)
            doc.add_paragraph()


def _render_footer(doc, footer):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"{footer.get('data_source', 'Data source: Apify Instagram Scraper')} | "
        f"{footer.get('posts_scraped', 0)} posts across {footer.get('accounts_count', 0)} accounts | "
        f"Scraped {footer.get('scrape_period', '')}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = INK_MUTED

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(
        f"{footer.get('analysis_powered_by', 'Analysis powered by Claude AI')} | "
        f"{footer.get('prepared_for', '')}"
    )
    run2.font.size = Pt(9)
    run2.font.color.rgb = INK_MUTED


# ── Public entrypoint ────────────────────────────────────────────────────────

def build_docx_bytes(report: dict) -> bytes:
    doc = Document()
    _configure_styles(doc)

    _render_header(doc, report)
    _render_section1(doc, report.get("section1", {}) or {})
    _render_section2(doc, report.get("section2", {}) or {})
    _render_section3(doc, report.get("section3", {}) or {})
    _render_section4(doc, report.get("section4", {}) or {})
    _render_section5(doc, report.get("section5", {}) or {})
    _render_section6(doc, report.get("section6", {}) or {})
    _render_summary(doc, report.get("summary", {}) or {})
    _render_footer(doc, report.get("footer", {}) or {})

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
